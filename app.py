import io
import json
import os
import re
import uuid
from datetime import date

from flask import Flask, render_template, request, send_file, jsonify
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

app = Flask(__name__)
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Students-info sheet layout:
# Row 8 = headers (A8=Major, B8=Student First Name, C8=Student Last Name, D8=ID)
# Rows 9-59 = student data (up to 51 students)
STUDENT_DATA_START_ROW = 9
STUDENT_DATA_END_ROW = 59
COL_MAJOR = 1       # A
COL_FIRST_NAME = 2  # B
COL_LAST_NAME = 3   # C
COL_ID = 4          # D


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    student_file = request.files.get("student_list")
    gradebook_file = request.files.get("gradebook")

    if not student_file or not gradebook_file:
        return jsonify({"error": "Both files are required."}), 400

    try:
        # --- Read the student list from "Students Grade" sheet ---
        student_wb = load_workbook(student_file, read_only=True, data_only=True)

        student_ws = None
        for name in student_wb.sheetnames:
            if name.lower().strip() == "students grade":
                student_ws = student_wb[name]
                break

        if student_ws is None:
            available = ", ".join(student_wb.sheetnames)
            student_wb.close()
            return jsonify({"error": f'"Students Grade" sheet not found. Available sheets: {available}'}), 400

        rows = list(student_ws.iter_rows(values_only=True))
        student_wb.close()

        if len(rows) < 4:
            return jsonify({"error": "Students Grade sheet has no data rows."}), 400

        # Find the header row dynamically by looking for "STUDENT ID"
        header_idx = None
        id_col = None
        name_col = None
        major_col = None
        for ri, row in enumerate(rows[:5]):
            for ci, cell in enumerate(row):
                if cell is not None:
                    val = str(cell).strip().upper()
                    if val == "STUDENT ID":
                        id_col = ci
                        header_idx = ri
                    elif val == "STUDENT NAME":
                        name_col = ci
                    elif val == "MAJOR":
                        major_col = ci
            if header_idx is not None:
                break

        if header_idx is None or id_col is None or name_col is None:
            return jsonify({"error": "Could not find STUDENT ID / STUDENT NAME columns in the Students Grade sheet."}), 400

        # Find the first data row (skip sub-headers)
        data_start = header_idx + 1
        for ri in range(header_idx + 1, len(rows)):
            row = rows[ri]
            cell_val = row[id_col] if id_col < len(row) else None
            if cell_val is not None:
                val = str(cell_val).strip()
                if val and val.replace(" ", "").replace("\n", "") not in ("", "SCOREDOUTOF") and any(c.isdigit() for c in val):
                    data_start = ri
                    break

        # Extract student data
        students = []
        for row in rows[data_start:]:
            sid = row[id_col] if id_col < len(row) else None
            full_name = row[name_col] if name_col < len(row) else None
            major = row[major_col] if major_col is not None and major_col < len(row) else None

            if sid is None or full_name is None:
                continue

            sid = str(sid).strip()
            full_name = str(full_name).strip()

            if not sid or not full_name:
                continue

            # Split full name: first word = first name, rest = last name
            name_parts = full_name.split()
            first = name_parts[0] if name_parts else ""
            last = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

            major_val = str(major).strip() if major else ""

            students.append({"first": first, "last": last, "id": sid, "major": major_val})

        if not students:
            return jsonify({"error": "No student data found."}), 400

        # --- Load the gradebook .xlsm template (preserve VBA macros) ---
        gradebook_wb = load_workbook(gradebook_file, keep_vba=True)

        # Find Students-info sheet
        target_ws = None
        for name in gradebook_wb.sheetnames:
            if name.lower().strip() == "students-info":
                target_ws = gradebook_wb[name]
                break

        if target_ws is None:
            available = ", ".join(gradebook_wb.sheetnames)
            gradebook_wb.close()
            return jsonify({"error": f'"Students-info" sheet not found. Available sheets: {available}'}), 400

        # Clear existing student data in rows 9-59
        for r in range(STUDENT_DATA_START_ROW, STUDENT_DATA_END_ROW + 1):
            target_ws.cell(row=r, column=COL_MAJOR).value = None
            target_ws.cell(row=r, column=COL_FIRST_NAME).value = None
            target_ws.cell(row=r, column=COL_LAST_NAME).value = None
            target_ws.cell(row=r, column=COL_ID).value = None

        # Write student data starting at row 9
        max_students = STUDENT_DATA_END_ROW - STUDENT_DATA_START_ROW + 1  # 51
        written = 0
        for i, s in enumerate(students[:max_students]):
            r = STUDENT_DATA_START_ROW + i
            if s["major"]:
                target_ws.cell(row=r, column=COL_MAJOR, value=s["major"])
            target_ws.cell(row=r, column=COL_FIRST_NAME, value=s["first"])
            target_ws.cell(row=r, column=COL_LAST_NAME, value=s["last"])
            target_ws.cell(row=r, column=COL_ID, value=s["id"])
            written += 1

        # Save as .xlsm (preserving macros)
        unique_id = uuid.uuid4().hex[:8]
        original_name = gradebook_file.filename or "gradebook.xlsm"
        safe_name = f"{unique_id}_{original_name}"
        output_path = os.path.join(UPLOAD_FOLDER, safe_name)
        gradebook_wb.save(output_path)
        gradebook_wb.close()

        # Build preview (first 10 students)
        preview = []
        for s in students[:10]:
            preview.append({
                "first": s["first"],
                "last": s["last"],
                "id": str(s["id"]),
                "major": s["major"],
            })

        return jsonify({
            "success": True,
            "message": f"Successfully filled {written} students into the gradebook.",
            "filename": safe_name,
            "original_name": original_name,
            "preview": preview,
            "total_students": len(students),
            "written": written,
            "skipped": max(0, len(students) - max_students),
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/headers", methods=["POST"])
def get_headers():
    """Read headers from the 'Students Grade' sheet and return student preview info."""
    student_file = request.files.get("student_list")
    if not student_file:
        return jsonify({"error": "No file provided."}), 400

    try:
        wb = load_workbook(student_file, read_only=True, data_only=True)

        ws = None
        for name in wb.sheetnames:
            if name.lower().strip() == "students grade":
                ws = wb[name]
                break

        if ws is None:
            available = ", ".join(wb.sheetnames)
            wb.close()
            return jsonify({"error": f'"Students Grade" sheet not found. Available sheets: {available}'}), 400

        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        headers = ["First Name (auto)", "Last Name (auto)", "Student ID", "Major"]
        guesses = {
            "first_name": "First Name (auto)",
            "last_name": "Last Name (auto)",
            "id": "Student ID",
            "major": "Major",
        }

        return jsonify({"headers": headers, "guesses": guesses})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/download/<filename>")
def download(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.isfile(file_path):
        return "File not found.", 404

    # Extract original name (strip the uuid prefix)
    original_name = filename.split("_", 1)[1] if "_" in filename else filename

    # Determine MIME type based on extension
    if original_name.endswith(".xlsm"):
        mimetype = "application/vnd.ms-excel.sheet.macroEnabled.12"
    elif original_name.endswith(".xlsx"):
        mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        mimetype = "application/octet-stream"

    response = send_file(
        file_path,
        as_attachment=True,
        download_name=original_name,
        mimetype=mimetype,
    )
    response.call_on_close(lambda: cleanup_file(file_path))
    return response


def cleanup_file(path):
    try:
        os.remove(path)
    except OSError:
        pass


FACULTY_FILE = os.path.join(UPLOAD_FOLDER, "faculty_assignments.xlsx")

DAY_MAP = {"U": "Sunday", "M": "Monday", "T": "Tuesday", "W": "Wednesday", "R": "Thursday"}
DAY_ORDER = ["M", "T", "U", "W", "R"]

# Column header names to search for dynamically
FACULTY_HEADERS = {
    "faculty": "faculty assignment",
    "course": "course",
    "type": "type",
    "section": "sec",
    "crn": "crn",
    "time": "time",
    "ramadan_time": "ramadan_time",
    "days": "days",
    "room": "room",
    "unit": "faculty unit",
}


def find_faculty_columns(ws):
    """Read the header row and return a dict mapping logical names to column indices."""
    cols = {}
    for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
        for ci, cell in enumerate(row):
            if cell is None:
                continue
            val = str(cell).strip().lower()
            for key, header in FACULTY_HEADERS.items():
                if key not in cols and val == header:
                    cols[key] = ci
        break
    # Fallback: match partial for section (could be "Sec" or "Section")
    if "section" not in cols:
        for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
            for ci, cell in enumerate(row):
                if cell and "sec" in str(cell).strip().lower():
                    cols["section"] = ci
                    break
            break
    return cols


@app.route("/faculty")
def faculty():
    return render_template("faculty.html")


@app.route("/faculty/status")
def faculty_status():
    """Check if a faculty file already exists and return its names."""
    if not os.path.isfile(FACULTY_FILE):
        return jsonify({"exists": False})

    try:
        wb, ws, cols = _open_faculty_ws()
        if wb is None:
            return jsonify({"exists": False})

        fc = cols.get("faculty")
        if fc is None:
            wb.close()
            return jsonify({"exists": False})

        names = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) > fc and row[fc]:
                name = str(row[fc]).strip()
                if name:
                    names.add(name)

        wb.close()
        return jsonify({"exists": True, "names": sorted(names)})
    except Exception:
        return jsonify({"exists": False})


@app.route("/faculty/upload", methods=["POST"])
def faculty_upload():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file provided."}), 400

    try:
        f.save(FACULTY_FILE)

        wb = load_workbook(FACULTY_FILE, read_only=True, data_only=True)
        ws = None
        for name in wb.sheetnames:
            if "faculty" in name.lower() and "assign" in name.lower():
                ws = wb[name]
                break
        if ws is None:
            ws = wb[wb.sheetnames[0]]

        cols = find_faculty_columns(ws)
        if "faculty" not in cols:
            wb.close()
            return jsonify({"error": "Could not find 'Faculty Assignment' column in the header row."}), 400

        fc = cols["faculty"]
        names = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) > fc and row[fc]:
                name = str(row[fc]).strip()
                if name:
                    names.add(name)

        wb.close()
        return jsonify({"success": True, "names": sorted(names)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/faculty/data")
def faculty_data():
    name = request.args.get("name", "").strip()
    if not name:
        return jsonify({"error": "No faculty name provided."}), 400

    if not os.path.isfile(FACULTY_FILE):
        return jsonify({"error": "No faculty file uploaded yet."}), 400

    try:
        wb = load_workbook(FACULTY_FILE, read_only=True, data_only=True)
        ws = None
        for sname in wb.sheetnames:
            if "faculty" in sname.lower() and "assign" in sname.lower():
                ws = wb[sname]
                break
        if ws is None:
            ws = wb[wb.sheetnames[0]]

        cols = find_faculty_columns(ws)
        if "faculty" not in cols:
            wb.close()
            return jsonify({"error": "Could not find 'Faculty Assignment' column in the header row."}), 400

        def cell(row, key):
            idx = cols.get(key)
            if idx is not None and idx < len(row) and row[idx] is not None:
                return str(row[idx]).strip()
            return ""

        schedule = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            faculty_name = cell(row, "faculty")
            if faculty_name.lower() != name.lower():
                continue

            course = cell(row, "course")
            ctype = cell(row, "type")
            section = cell(row, "section")
            crn = cell(row, "crn")
            time_raw = cell(row, "time")
            days = cell(row, "days")
            room = cell(row, "room")

            time_start = ""
            time_end = ""
            if "::" in time_raw:
                parts = time_raw.split("::")
                if len(parts) == 2:
                    s, e = parts[0].strip(), parts[1].strip()
                    if len(s) == 4 and s.isdigit():
                        time_start = s[:2] + ":" + s[2:]
                    if len(e) == 4 and e.isdigit():
                        time_end = e[:2] + ":" + e[2:]

            ramadan_time_raw = cell(row, "ramadan_time")
            # Discard formula errors like #N/A
            if ramadan_time_raw.startswith("#"):
                ramadan_time_raw = ""
            ramadan_time_start = ""
            ramadan_time_end = ""
            if "::" in ramadan_time_raw:
                rparts = ramadan_time_raw.split("::")
                if len(rparts) == 2:
                    rs, re_ = rparts[0].strip(), rparts[1].strip()
                    if len(rs) == 4 and rs.isdigit():
                        ramadan_time_start = rs[:2] + ":" + rs[2:]
                    if len(re_) == 4 and re_.isdigit():
                        ramadan_time_end = re_[:2] + ":" + re_[2:]

            day_names = [DAY_MAP[d] for d in days if d in DAY_MAP]

            schedule.append({
                "course": course,
                "type": ctype,
                "section": section,
                "crn": crn,
                "time_raw": time_raw,
                "time_start": time_start,
                "time_end": time_end,
                "ramadan_time_raw": ramadan_time_raw,
                "ramadan_time_start": ramadan_time_start,
                "ramadan_time_end": ramadan_time_end,
                "days": days,
                "day_names": day_names,
                "room": room,
            })

        wb.close()
        return jsonify({"faculty": name, "schedule": schedule})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _open_faculty_ws():
    """Helper: open the faculty file and return (wb, ws, cols) or raise."""
    if not os.path.isfile(FACULTY_FILE):
        return None, None, None
    wb = load_workbook(FACULTY_FILE, read_only=True, data_only=True)
    ws = None
    for sname in wb.sheetnames:
        if "faculty" in sname.lower() and "assign" in sname.lower():
            ws = wb[sname]
            break
    if ws is None:
        ws = wb[wb.sheetnames[0]]
    cols = find_faculty_columns(ws)
    return wb, ws, cols


@app.route("/faculty/programs")
def faculty_programs():
    wb, ws, cols = _open_faculty_ws()
    if wb is None:
        return jsonify({"error": "No faculty file uploaded yet."}), 400

    try:
        uc = cols.get("unit")
        if uc is None:
            wb.close()
            return jsonify({"error": "Could not find 'Faculty unit' column."}), 400

        programs = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) > uc and row[uc]:
                val = str(row[uc]).strip()
                if val:
                    programs.add(val)

        wb.close()
        return jsonify({"programs": sorted(programs)})
    except Exception as e:
        wb.close()
        return jsonify({"error": str(e)}), 500


SHIFT_DURATION = 8 * 60       # 8 hours in minutes
SHIFT_EARLY_BUFFER = 30       # faculty arrives 30 min before first class


def _parse_time(raw):
    """Parse 'HHMM::HHMM' into (start_minutes, end_minutes) from midnight."""
    if "::" not in raw:
        return None, None
    parts = raw.split("::")
    if len(parts) != 2:
        return None, None
    s, e = parts[0].strip(), parts[1].strip()
    if len(s) == 4 and s.isdigit() and len(e) == 4 and e.isdigit():
        return int(s[:2]) * 60 + int(s[2:]), int(e[:2]) * 60 + int(e[2:])
    return None, None


@app.route("/faculty/by-program")
def faculty_by_program():
    program = request.args.get("program", "").strip()
    if not program:
        return jsonify({"error": "No program provided."}), 400

    # Optional filters: only return faculty free at this time/day
    filter_time = request.args.get("time", "").strip()   # e.g. "0800::0850"
    filter_days = request.args.get("days", "").strip()    # e.g. "UT"
    time_mode = request.args.get("time_mode", "normal").strip()  # "normal" or "ramadan"

    wb, ws, cols = _open_faculty_ws()
    if wb is None:
        return jsonify({"error": "No faculty file uploaded yet."}), 400

    try:
        fc = cols.get("faculty")
        uc = cols.get("unit")
        tc = cols.get("ramadan_time") if time_mode == "ramadan" else cols.get("time")
        dc = cols.get("days")
        if fc is None or uc is None:
            wb.close()
            return jsonify({"error": "Required columns not found."}), 400

        def cell(row, idx):
            if idx is not None and idx < len(row) and row[idx] is not None:
                return str(row[idx]).strip()
            return ""

        # Collect all faculty in this program
        program_faculty = set()
        # Per-faculty per-day class times: {name: {day_letter: [(start_min, end_min), ...]}}
        faculty_classes = {}

        for row in ws.iter_rows(min_row=2, values_only=True):
            unit_val = cell(row, uc)
            fac_val = cell(row, fc)
            if not fac_val:
                continue
            if unit_val.lower() == program.lower():
                program_faculty.add(fac_val)

        # Scan ALL rows to get full schedule for program faculty
        for row in ws.iter_rows(min_row=2, values_only=True):
            fac_val = cell(row, fc)
            if fac_val not in program_faculty:
                continue
            row_time = cell(row, tc)
            row_days = cell(row, dc)
            t_start, t_end = _parse_time(row_time)
            if t_start is None:
                continue
            if fac_val not in faculty_classes:
                faculty_classes[fac_val] = {}
            for d in row_days:
                if d not in faculty_classes[fac_val]:
                    faculty_classes[fac_val][d] = []
                faculty_classes[fac_val][d].append((t_start, t_end))

        busy_faculty = set()
        session_start, session_end = _parse_time(filter_time) if filter_time else (None, None)

        if session_start is not None and filter_days:
            for fac_name in program_faculty:
                classes = faculty_classes.get(fac_name, {})
                for d in filter_days:
                    day_classes = classes.get(d, [])

                    # Direct conflict: already teaching at this exact time
                    for cs, ce in day_classes:
                        if cs < session_end and ce > session_start:
                            busy_faculty.add(fac_name)
                            break

                    if fac_name in busy_faculty:
                        break

                    # Shift check: if faculty has classes on this day,
                    # calculate their 8-hour shift window
                    if day_classes:
                        earliest_start = min(cs for cs, ce in day_classes)
                        latest_end = max(ce for cs, ce in day_classes)

                        # Shift starts 30 min before earliest class
                        shift_start = earliest_start - SHIFT_EARLY_BUFFER
                        shift_end = shift_start + SHIFT_DURATION

                        # If latest class pushes past shift_end, anchor to latest end
                        if latest_end > shift_end:
                            shift_end = latest_end
                            shift_start = shift_end - SHIFT_DURATION

                        # Session must fit within shift window
                        if session_start < shift_start or session_end > shift_end:
                            busy_faculty.add(fac_name)
                            break

        available = sorted(program_faculty - busy_faculty)

        wb.close()
        return jsonify({"faculty": available, "busy": sorted(busy_faculty)})
    except Exception as e:
        wb.close()
        return jsonify({"error": str(e)}), 500


RECORDS_FILE = os.path.join(UPLOAD_FOLDER, "team_records.json")


def _load_records():
    if not os.path.isfile(RECORDS_FILE):
        return []
    with open(RECORDS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("records", [])


def _save_records(records):
    with open(RECORDS_FILE, "w", encoding="utf-8") as f:
        json.dump({"records": records}, f, ensure_ascii=False, indent=2)


@app.route("/records")
def records_page():
    return render_template("records.html")


@app.route("/records/data")
def records_data():
    return jsonify({"records": _load_records()})


@app.route("/records/add", methods=["POST"])
def records_add():
    body = request.get_json(force=True)
    name = (body.get("name") or "").strip()
    rtype = (body.get("type") or "").strip()
    description = (body.get("description") or "").strip()
    rdate = (body.get("date") or "").strip() or date.today().isoformat()
    semester = (body.get("semester") or "").strip()
    year = (body.get("year") or "").strip()

    if not name or not rtype or not description:
        return jsonify({"error": "Name, type, and description are required."}), 400
    if rtype not in ("achievement", "refusal", "note"):
        return jsonify({"error": "Invalid type."}), 400

    records = _load_records()
    record = {
        "id": uuid.uuid4().hex[:8],
        "name": name,
        "type": rtype,
        "description": description,
        "date": rdate,
        "semester": semester,
        "year": year,
    }
    records.append(record)
    _save_records(records)
    return jsonify({"success": True, "record": record})


@app.route("/records/delete/<record_id>", methods=["DELETE"])
def records_delete(record_id):
    records = _load_records()
    new_records = [r for r in records if r["id"] != record_id]
    if len(new_records) == len(records):
        return jsonify({"error": "Record not found."}), 404
    _save_records(new_records)
    return jsonify({"success": True})


@app.route("/records/export")
def records_export():
    records = _load_records()

    wb = Workbook()
    ws = wb.active
    ws.title = "Team Records"

    # Styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="1D1D1F", end_color="1D1D1F", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="D2D2D7"),
        right=Side(style="thin", color="D2D2D7"),
        top=Side(style="thin", color="D2D2D7"),
        bottom=Side(style="thin", color="D2D2D7"),
    )
    type_fills = {
        "achievement": PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        "refusal": PatternFill(start_color="FECACA", end_color="FECACA", fill_type="solid"),
        "note": PatternFill(start_color="E8E8ED", end_color="E8E8ED", fill_type="solid"),
    }

    headers = ["Date", "Member", "Type", "Semester", "Year", "Description"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    for ri, rec in enumerate(records, 2):
        vals = [
            rec.get("date", ""),
            rec.get("name", ""),
            rec.get("type", ""),
            rec.get("semester", ""),
            rec.get("year", ""),
            rec.get("description", ""),
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws.cell(row=ri, column=ci, value=v)
            cell.border = thin_border
            if ci == 3:  # type column
                fill = type_fills.get(v)
                if fill:
                    cell.fill = fill

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 24
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 8
    ws.column_dimensions["F"].width = 50

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    wb.close()

    return send_file(
        buf,
        as_attachment=True,
        download_name="team_records.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


BONUS_FILE = os.path.join(UPLOAD_FOLDER, "bonus_data.xlsx")
BONUS_JSON = os.path.join(UPLOAD_FOLDER, "bonus_data.json")


@app.route("/bonus")
def bonus_page():
    return render_template("bonus.html")


@app.route("/bonus/upload", methods=["POST"])
def bonus_upload():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file provided."}), 400

    try:
        f.save(BONUS_FILE)

        wb = load_workbook(BONUS_FILE, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        if len(rows) < 3:
            return jsonify({"error": "File has too few rows."}), 400

        # Find header row dynamically (look for "Student ID" in first 5 rows)
        header_idx = None
        id_col = None
        name_col = None
        dept_col = None
        major_col = None
        for ri, row in enumerate(rows[:5]):
            for ci, cell in enumerate(row):
                if cell is not None:
                    val = str(cell).strip().upper()
                    if val == "STUDENT ID":
                        id_col = ci
                        header_idx = ri
                    elif val == "STUDENT NAME":
                        name_col = ci
                    elif "DEPARTMENT" in val:
                        dept_col = ci
                    elif "MAJOR" in val:
                        major_col = ci
            if header_idx is not None:
                break

        if header_idx is None or id_col is None or name_col is None:
            return jsonify({"error": "Could not find Student ID / Student Name columns."}), 400

        # Find bonus column: search title and header rows for "bonus" (case-insensitive)
        bonus_col = None
        for ri in range(min(header_idx + 1, len(rows))):
            for ci, cell in enumerate(rows[ri]):
                if cell is not None and "bonus" in str(cell).strip().lower():
                    bonus_col = ci
                    break
            if bonus_col is not None:
                break

        if bonus_col is None:
            return jsonify({"error": "Could not find 'Bonus' column."}), 400

        # Find course columns: headers matching "Course \d+" pattern
        header_row = rows[header_idx]
        course_cols = []
        for ci, cell in enumerate(header_row):
            if cell is not None and re.match(r"(?i)course\s*\d+", str(cell).strip()):
                course_cols.append(ci)

        # Parse students
        students = []
        for row in rows[header_idx + 1:]:
            sid = row[id_col] if id_col < len(row) else None
            sname = row[name_col] if name_col < len(row) else None
            if sid is None or sname is None:
                continue
            sid = str(sid).strip()
            sname = str(sname).strip()
            if not sid or not sname:
                continue

            dept = str(row[dept_col]).strip() if dept_col is not None and dept_col < len(row) and row[dept_col] else ""
            major = str(row[major_col]).strip() if major_col is not None and major_col < len(row) and row[major_col] else ""

            bonus_val = row[bonus_col] if bonus_col < len(row) else None
            try:
                bonus = float(bonus_val) if bonus_val is not None else 0
            except (ValueError, TypeError):
                bonus = 0

            courses = []
            for cc in course_cols:
                if cc < len(row) and row[cc]:
                    raw = str(row[cc]).strip()
                    if not raw:
                        continue
                    # Split on dash or space
                    if "-" in raw:
                        parts = raw.split("-", 1)
                    else:
                        parts = raw.rsplit(" ", 1)
                    course_name = parts[0].strip()
                    section = parts[1].strip() if len(parts) > 1 else ""
                    courses.append({"course": course_name, "section": section})

            students.append({
                "name": sname,
                "id": sid,
                "department": dept,
                "major": major,
                "courses": courses,
                "bonus": bonus,
            })

        if not students:
            return jsonify({"error": "No student data found."}), 400

        majors = sorted(set(s["major"] for s in students if s["major"]))

        with open(BONUS_JSON, "w", encoding="utf-8") as jf:
            json.dump({"majors": majors, "data": students}, jf, ensure_ascii=False, indent=2)

        return jsonify({"success": True, "majors": majors, "data": students})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/bonus/data")
def bonus_data():
    if not os.path.isfile(BONUS_JSON):
        return jsonify({"error": "No bonus data uploaded yet."}), 400
    with open(BONUS_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)
    return jsonify(data)


@app.route("/referral/generate", methods=["POST"])
def referral_generate():
    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        body = request.get_json(force=True)
        student = body.get("student", {})
        ci = body.get("course_info", {})

        template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Referral Form.docx")
        doc = DocxDocument(template_path)
        paras = doc.paragraphs

        # ── Parse term ────────────────────────────────────────────────────
        term_full = ci.get("term", "")
        semester = term_full
        academic_year = ""
        m = re.search(r"(Fall|Spring|Summer)\s+(\d{4})", term_full, re.IGNORECASE)
        if m:
            semester = m.group(1) + " " + m.group(2)
            yr = int(m.group(2))
            academic_year = f"{yr}/{yr+1}" if m.group(1).lower() == "fall" else f"{yr-1}/{yr}"

        course_code = ci.get("course_code", "")
        m2 = re.match(r"([A-Z]+-\d+)", course_code)
        if m2:
            course_code = m2.group(1)

        today_str = date.today().strftime("%d %B %Y")
        instructor         = student.get("instructor", "")
        section            = student.get("section", "")
        student_name       = student.get("name", "")
        student_id         = student.get("id", "")
        major              = student.get("major", "")
        gpa                = student.get("gpa", "")
        acd                = student.get("acd", "")
        after_avg          = student.get("after_avg", "")
        after_grd          = student.get("after_grd", "")
        lab_summary        = student.get("lab_summary", [])
        consecutive_issues = student.get("consecutive_issues", [])

        # ── Helpers ───────────────────────────────────────────────────────
        def clear_runs(para):
            for r in list(para._p.findall(qn("w:r"))):
                para._p.remove(r)

        def append_value(para, text):
            para.add_run("  " + text)

        def rewrite(para, text):
            clear_runs(para)
            para.add_run(text)

        def add_para(ref_p, text="", bold=False, indent=False):
            """Create a Body Text paragraph and insert it after ref_p (lxml element).
            Returns the new lxml element so it can be used as the next ref_p."""
            new_para = doc.add_paragraph(style="Body Text")
            new_para.paragraph_format.space_before = Pt(0)
            new_para.paragraph_format.space_after = Pt(1)
            if indent:
                new_para.paragraph_format.left_indent = Inches(0.3)
            if text:
                run = new_para.add_run(text)
                run.bold = bold
            # Move element from end of document to after ref_p
            new_p = new_para._p
            new_p.getparent().remove(new_p)
            ref_p.addnext(new_p)
            return new_p

        def lab_desc(l):
            if not l:
                return ""
            if l["status"] == "absent":
                return f"Lab {l['lab']} (Absent)"
            avg = f"{l['avg']}%" if l["avg"] is not None else "< 60%"
            return f"Lab {l['lab']} ({avg})"

        # ── Fill header fields ────────────────────────────────────────────
        append_value(paras[2], instructor)
        append_value(paras[4], semester)
        append_value(paras[7], course_code)
        rewrite(paras[8], f"Academic Year :  {academic_year}          Section :  {section}")
        rewrite(paras[13], "\u2611  Academic\t\t Non-Academic")

        # ── Build "Please Elaborate" section ─────────────────────────────
        # Insert paragraphs one by one after para[14], tracking last inserted element
        ref_p = paras[14]._p

        absent_labs = [l for l in lab_summary if l["status"] == "absent"]
        weak_labs   = [l for l in lab_summary if l["status"] == "weak"]

        # Student info
        ref_p = add_para(ref_p, f"Student Name  :   {student_name}")
        ref_p = add_para(ref_p, f"Student ID      :   {student_id}")
        ref_p = add_para(ref_p, f"Major              :   {major}")
        ref_p = add_para(ref_p, f"Current Grade :   {after_avg} ({after_grd})   |   GPA: {gpa}   |   Academic Status: {acd}")
        ref_p = add_para(ref_p)  # blank

        # Attendance section
        ref_p = add_para(ref_p, "Attendance Issues:", bold=True)
        if absent_labs:
            for l in absent_labs:
                ref_p = add_para(ref_p, f"Lab {l['lab']}  \u2014  Absent", indent=True)
        else:
            ref_p = add_para(ref_p, "No absent labs.", indent=True)
        ref_p = add_para(ref_p)  # blank

        # Weak performance section
        ref_p = add_para(ref_p, "Weak Performance Issues  (score < 60%):", bold=True)
        if weak_labs:
            for l in weak_labs:
                avg_str = f"{l['avg']}%" if l["avg"] is not None else "< 60%"
                ref_p = add_para(ref_p, f"Lab {l['lab']}  \u2014  {avg_str}", indent=True)
        else:
            ref_p = add_para(ref_p, "No weak performance labs.", indent=True)
        ref_p = add_para(ref_p)  # blank

        # Consecutive issues
        ref_p = add_para(ref_p, "Consecutive Issues Identified:", bold=True)
        for pair in consecutive_issues:
            l1 = next((l for l in lab_summary if l["lab"] == pair[0]), None)
            l2 = next((l for l in lab_summary if l["lab"] == pair[1]), None)
            both_absent = l1 and l2 and l1["status"] == "absent" and l2["status"] == "absent"
            action = "Missed" if both_absent else "Underperformed in"
            ref_p = add_para(ref_p, f"\u2022  {action}:  {lab_desc(l1)}  and  {lab_desc(l2)}", indent=True)
        ref_p = add_para(ref_p)  # blank

        # Full lab summary
        ref_p = add_para(ref_p, "Full Lab Performance Summary:", bold=True)
        for l in lab_summary:
            if l["status"] == "absent":
                st = "Absent"
            elif l["status"] == "weak":
                st = f"{l['avg']}%  (Weak)" if l["avg"] is not None else "Weak (< 60%)"
            elif l["status"] == "ok":
                st = f"{l['avg']}%  (OK)" if l["avg"] is not None else "OK"
            else:
                st = "N/A"
            ref_p = add_para(ref_p, f"Lab {l['lab']}  :  {st}", indent=True)

        # ── Signature line ────────────────────────────────────────────────
        rewrite(paras[17], f"Instructor's Name & Signature :  {instructor}\t\tDate:  {today_str}")

        # ── Return docx ───────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_name = re.sub(r"[^A-Za-z0-9_\-]", "_", student_name)
        filename = f"Referral_{student_id}_{safe_name}.docx"
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/referral")
def referral_page():
    return render_template("referral.html")


@app.route("/referral/parse", methods=["POST"])
def referral_parse():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file provided."}), 400

    try:
        wb = load_workbook(f, read_only=True, data_only=True)

        # ── Cover Page info ──────────────────────────────────────────────
        course_info = {}
        if "Cover Page" in wb.sheetnames:
            cp_rows = list(wb["Cover Page"].iter_rows(values_only=True))
            for row in cp_rows:
                for i, cell in enumerate(row):
                    if cell is None:
                        continue
                    key = str(cell).strip().upper()
                    val = str(row[i + 1]).strip() if i + 1 < len(row) and row[i + 1] is not None else ""
                    if key == "TERM":
                        course_info["term"] = val
                    elif key == "COURSE CODE":
                        course_info["course_code"] = val
                    elif key == "COURSE TITLE":
                        course_info["course_title"] = val
                    elif key == "HOD":
                        course_info["hod"] = val
                    elif key == "DEAN":
                        course_info["dean"] = val
                    elif key == "NO. OF STUDENTS":
                        course_info["n_students"] = val

        # ── Students Grade sheet ─────────────────────────────────────────
        ws = None
        for sname in wb.sheetnames:
            if sname.lower().strip() == "students grade":
                ws = wb[sname]
                break
        if ws is None:
            wb.close()
            return jsonify({"error": '"Students Grade" sheet not found.'}), 400

        rows = list(ws.iter_rows(values_only=True))

        # Find header row
        header_idx = None
        for ri, row in enumerate(rows[:5]):
            for ci, cell in enumerate(row):
                if cell is not None and str(cell).strip().upper() == "STUDENT ID":
                    header_idx = ri
                    break
            if header_idx is not None:
                break
        if header_idx is None:
            wb.close()
            return jsonify({"error": "Could not find header row."}), 400

        hrow = rows[header_idx]
        sub_row = rows[header_idx + 1] if header_idx + 1 < len(rows) else []

        id_col = name_col = major_col = section_col = instructor_col = crn_col = None
        gpa_col = acd_col = after_bonus_avg_col = after_bonus_grd_col = None

        for ci, cell in enumerate(hrow):
            if cell is None:
                continue
            val = str(cell).strip().upper().replace("\n", " ").strip()
            if val == "STUDENT ID":
                id_col = ci
            elif val == "STUDENT NAME":
                name_col = ci
            elif val == "MAJOR":
                major_col = ci
            elif val == "CRN":
                crn_col = ci
            elif val == "SECTION":
                section_col = ci
            elif val == "INSTRUCTOR":
                instructor_col = ci
            elif "GPA" in val:
                gpa_col = ci
            elif "ACD" in val and "STD" in val:
                acd_col = ci

        for ci, cell in enumerate(hrow):
            if cell is None:
                continue
            val = str(cell).strip().upper()
            sub = str(sub_row[ci]).strip().upper() if ci < len(sub_row) and sub_row[ci] is not None else ""
            if val == "AFTER BONUS":
                if sub == "AVG" and after_bonus_avg_col is None:
                    after_bonus_avg_col = ci
                elif sub == "GRD" and after_bonus_grd_col is None:
                    after_bonus_grd_col = ci

        # ── Parse lab columns ─────────────────────────────────────────────
        lab_data = {}  # {lab_num: {exp_avg: ci, rep_avg: ci, exp_name: str}}
        for ci, cell in enumerate(hrow):
            if cell is None:
                continue
            cell_str = str(cell).strip()
            m = re.match(r"Lab\s+(\d+)", cell_str, re.IGNORECASE)
            if not m:
                continue
            lab_num = int(m.group(1))
            sub = str(sub_row[ci]).strip().upper() if ci < len(sub_row) and sub_row[ci] is not None else ""
            if sub != "AVG":
                continue
            if lab_num not in lab_data:
                lab_data[lab_num] = {"exp_avg": None, "rep_avg": None, "exp_name": cell_str, "rep_name": ""}
            is_exp = bool(re.search(r"experiment|exp\b", cell_str, re.IGNORECASE))
            is_rep = "report" in cell_str.lower()
            if is_exp:
                lab_data[lab_num]["exp_avg"] = ci
                lab_data[lab_num]["exp_name"] = cell_str
            elif is_rep:
                lab_data[lab_num]["rep_avg"] = ci
                lab_data[lab_num]["rep_name"] = cell_str

        sorted_labs = sorted(lab_data.keys())
        data_start = header_idx + 2

        def get_lab_status(row, lab_num):
            info = lab_data[lab_num]
            avg_col = info.get("exp_avg") or info.get("rep_avg")
            if avg_col is None or avg_col >= len(row):
                return "unknown", None
            val = row[avg_col]
            if val is None:
                return "unknown", None
            val_str = str(val).strip()
            if "ABSENT" in val_str.upper():
                return "absent", None
            try:
                num = float(val)
                if num < 60:
                    return "weak", round(num, 1)
                return "ok", round(num, 1)
            except (ValueError, TypeError):
                return "unknown", None

        flagged = []
        for row in rows[data_start:]:
            if id_col is None or id_col >= len(row):
                continue
            sid = row[id_col]
            if sid is None:
                continue
            sid = str(sid).strip()
            if not sid or not any(c.isdigit() for c in sid):
                continue

            def cv(col):
                if col is None or col >= len(row) or row[col] is None:
                    return ""
                return str(row[col]).strip()

            name = cv(name_col)
            major = cv(major_col)
            section = cv(section_col)
            instructor = cv(instructor_col)
            crn = cv(crn_col)
            gpa = cv(gpa_col)
            acd = cv(acd_col)
            after_avg = cv(after_bonus_avg_col)
            after_grd = cv(after_bonus_grd_col)

            lab_statuses = {}
            for n in sorted_labs:
                status, avg_num = get_lab_status(row, n)
                lab_statuses[n] = {"status": status, "avg": avg_num}

            # Check for 2 consecutive problematic labs
            consecutive_issues = []
            for i in range(len(sorted_labs) - 1):
                l1, l2 = sorted_labs[i], sorted_labs[i + 1]
                s1, s2 = lab_statuses[l1]["status"], lab_statuses[l2]["status"]
                if s1 in ("absent", "weak") and s2 in ("absent", "weak"):
                    already = any(l1 in p for p in consecutive_issues)
                    if not already:
                        consecutive_issues.append([l1, l2])

            if not consecutive_issues:
                continue

            lab_summary = []
            for n in sorted_labs:
                info = lab_statuses[n]
                lab_summary.append({
                    "lab": n,
                    "status": info["status"],
                    "avg": info["avg"],
                })

            flagged.append({
                "id": sid,
                "name": name,
                "major": major,
                "section": section,
                "instructor": instructor,
                "crn": crn,
                "gpa": gpa,
                "acd": acd,
                "after_avg": after_avg,
                "after_grd": after_grd,
                "consecutive_issues": consecutive_issues,
                "lab_summary": lab_summary,
            })

        wb.close()
        return jsonify({
            "success": True,
            "course_info": course_info,
            "flagged": flagged,
            "total_flagged": len(flagged),
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/meeting")
def meeting_page():
    return render_template("meeting.html")


@app.route("/meeting/generate", methods=["POST"])
def meeting_generate():
    try:
        from docx import Document as DocxDocument
        from docx.oxml import OxmlElement
        from copy import deepcopy
        from datetime import datetime

        data = request.get_json(force=True)

        date_str     = data.get("date", "")
        meeting_num  = int(data.get("number") or 1)
        facilitator  = (data.get("facilitator") or "").strip()
        time_val     = (data.get("time") or "TBD").strip() or "TBD"
        duration_val = str(data.get("duration") or "45").strip() or "45"
        location     = (data.get("location") or "Meeting Room").strip()
        attendees    = data.get("attendees") or []
        agenda       = data.get("agenda") or []
        apologies    = (data.get("apologies") or "None").strip() or "None"
        next_mtg     = (data.get("next_meeting") or "TBD").strip() or "TBD"

        try:
            dt  = datetime.strptime(date_str, "%Y-%m-%d")
            d   = dt.day
            sfx = "th" if 11 <= d <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(d % 10, "th")
            date_long = dt.strftime(f"%A, %B {d:02d}{sfx}, %Y")
            year = str(dt.year)
        except Exception:
            date_long = date_str or "TBD"
            year = "2026"

        att_names = [a.get("name", "").strip() for a in attendees if a.get("name", "").strip()]
        att_depts = [a.get("dept", "Computer Engineering Technology").strip() for a in attendees]
        if not att_names and facilitator:
            att_names = [facilitator]
            att_depts = ["Computer Engineering Technology"]
        atts_str = ", ".join(att_names)

        tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)), "1. Meeting03022026.docx")
        doc = DocxDocument(tpl)

        def cell_txt(tbl, row, col, text):
            cell = tbl.cell(row, col)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            p0 = cell.paragraphs[0]
            if p0.runs:
                p0.runs[0].text = text
            else:
                p0.add_run(text)

        def para_txt(para, text):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)

        def clone_with_text(src_para, text):
            from docx.oxml.ns import qn as _qn
            new_p = deepcopy(src_para._p)
            for t_el in new_p.findall(f".//{_qn('w:t')}"):
                t_el.text = ""
            first_t = new_p.find(f".//{_qn('w:t')}")
            if first_t is not None:
                first_t.text = text
            else:
                r_el = OxmlElement("w:r")
                t_el = OxmlElement("w:t")
                t_el.text = text
                r_el.append(t_el)
                new_p.append(r_el)
            return new_p

        # Capture table refs BEFORE any XML insertions so indices don’t shift
        t3 = doc.tables[3]
        t4 = doc.tables[4]
        t5 = doc.tables[5]

        def clone_t3_with_decision(src_tbl, decision_text):
            from docx.oxml.ns import qn as _qn
            WTR = _qn("w:tr")
            WTC = _qn("w:tc")
            WT  = _qn("w:t")
            new_tbl = deepcopy(src_tbl._tbl)
            rows = new_tbl.findall(WTR)
            if len(rows) > 1:
                tcs = rows[1].findall(WTC)
                if tcs:
                    tc = tcs[0]
                    for t_el in tc.findall(".//" + WT):
                        t_el.text = ""
                    first_t = tc.find(".//" + WT)
                    if first_t is not None:
                        first_t.text = decision_text
            return new_tbl

        # Table 0 – main header
        t0 = doc.tables[0]
        cell_txt(t0, 6, 0, f"Date of Meeting     :  {date_long}")
        cell_txt(t0, 6, 3, f"Duration of Meeting :   {duration_val}")
        cell_txt(t0, 7, 0, f"Meeting Facilitator: {facilitator}")
        cell_txt(t0, 7, 3, f"Location                       :   {location}")
        for i in range(7):
            if i < len(att_names):
                cell_txt(t0, 10 + i, 0, att_names[i])
                cell_txt(t0, 10 + i, 3, att_depts[i] if i < len(att_depts) else "Computer Engineering Technology")
            else:
                cell_txt(t0, 10 + i, 0, "")
                cell_txt(t0, 10 + i, 3, "")
        for i in range(6):
            cell_txt(t0, 20 + i, 0, agenda[i].get("topic", "") if i < len(agenda) else "")

        # Table 1 – meeting info block
        t1 = doc.tables[1]
        cell_txt(t1, 0, 1, facilitator)
        cell_txt(t1, 1, 1, date_long)
        cell_txt(t1, 2, 1, time_val)
        cell_txt(t1, 3, 1, location)
        cell_txt(t1, 4, 1, atts_str)
        cell_txt(t1, 5, 1, apologies)

        paras = doc.paragraphs
        para_txt(paras[4], f"Meeting # {meeting_num}")

        for i in range(7):
            if i < len(att_names):
                para_txt(paras[10 + i], f"{att_names[i]} moves to accept minutes of the meeting")
            else:
                para_txt(paras[10 + i], "")

        # First agenda item – reuse existing template paragraphs 21-27
        if agenda:
            a0 = agenda[0]
            a0_decs = (a0.get("decisions") or "").strip()
            para_txt(paras[21], "Agenda’s Item 1: ")
            para_txt(paras[22], "Proposed points: ")
            para_txt(paras[23], a0.get("topic", ""))
            para_txt(paras[24], "Decisions: ")
            dec_lines = a0_decs.split("\n")
            para_txt(paras[25], dec_lines[0] if dec_lines else "")
            para_txt(paras[26], dec_lines[1] if len(dec_lines) > 1 else "")
            para_txt(paras[27], "Actions: ")
            cell_txt(t3, 1, 0, a0_decs)
        else:
            for i in range(21, 28):
                para_txt(paras[i], "")

        # Additional agenda items – back-to-back (no blank separator), each with its own Actions table
        if len(agenda) > 1:
            anchor = paras[28]._p
            for idx in range(1, len(agenda)):
                ai    = agenda[idx]
                topic = ai.get("topic", "")
                decs  = (ai.get("decisions") or "").strip()
                for src_p, txt in [
                    (paras[21], f"Agenda’s Item {idx + 1}: "),
                    (paras[22], "Proposed points: "),
                    (paras[23], topic),
                    (paras[24], "Decisions: "),
                    (paras[25], decs),
                    (paras[27], "Actions: "),
                ]:
                    anchor.addprevious(clone_with_text(src_p, txt))
                anchor.addprevious(clone_t3_with_decision(t3, decs))

        # Table 2 – "committee met to review" list: topics go in col 1, rows 0-5
        t2 = doc.tables[2]
        for i in range(6):
            cell_txt(t2, i, 1, agenda[i].get("topic", "") if i < len(agenda) else "")

        # Table 4 – all decisions summary (rows 1-6)
        for i, ai in enumerate(agenda[:6]):
            topic = ai.get("topic", "")
            decs  = (ai.get("decisions") or "").strip()
            cell_txt(t4, i + 1, 0, f"Item {i+1} – {topic}: {decs}" if (topic or decs) else "")

        # Next-meeting paragraph (search by content so index shifts don’t matter)
        for p in doc.paragraphs:
            if "Next meeting" in p.text:
                para_txt(p, f"Next meeting  is scheduled on :     {next_mtg}")
                break

        # Table 5 – signature sheet
        for i in range(7):
            ri = i * 2
            if i < len(att_names):
                cell_txt(t5, ri, 0, att_names[i])
                cell_txt(t5, ri, 2, f"____ /____ /{year}")
            else:
                cell_txt(t5, ri, 0, "")
                cell_txt(t5, ri, 1, "")
                cell_txt(t5, ri, 2, "")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        compact = date_str.replace("-", "") if date_str else "unknown"
        fname = f"Meeting_{meeting_num}_{compact}.docx"
        return send_file(
            buf, as_attachment=True, download_name=fname,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


if __name__ == "__main__":
    app.run(debug=False)
